"""
selfloop_engine.py
==================
Streamlit 앱(app_selfloop.py)이 import 하는 통합 엔진.

구성:
  - EmbeddingProvider : tok_emb pkl 있으면 사용, 없으면 결정론적 해시 폴백
  - GrowingSOM        : 자율 성장 SOM (+ INT8 양자화 저장)
  - Crawler           : "이 분야 검색해서 학습" 명령 -> 웹 수집 (네트워크 필요)
  - LLMBridge         : 답변 루프용 LLM 호출 (API 키 필요)
  - SelfLoopState     : 전체 상태 + pkl 저장/불러오기(이전 학습 누적)
  - introspect 지표   : 점유율/QE/어휘다양성/추론경로

네트워크/키가 없는 환경에서도 코퍼스 학습·계측·pkl 은 완전 작동.
크롤링/LLM 은 로컬에서 키 넣으면 작동.
"""

from __future__ import annotations
import os
import re
import time
import pickle
import hashlib
import math
import numpy as np


# ======================================================================
# 임베딩
# ======================================================================
def tokenize(sentence: str):
    """한국어 어절 + 영문 단어 토크나이즈 (GrowingEmbedding/build_tok_emb 공용)."""
    s = sentence.strip()
    s = re.sub(r"https?://\S+|\S+@\S+", " ", s)
    toks = re.findall(r"[가-힣]+|[A-Za-z]+|[0-9]+", s)
    out = []
    for t in toks:
        if re.fullmatch(r"[A-Za-z0-9]", t):
            continue
        out.append(t.lower() if re.fullmatch(r"[A-Za-z]+", t) else t)
    return out


class EmbeddingProvider:
    """
    실제 TinyTransformer tok_emb를 받아 사용. 없거나 형식 불명이면 해시 폴백.

    인식하는 pkl 형식(자동 감지):
      1) {"word2idx": {...}, "tok_emb": ndarray(V, dim)}            # 기본
      2) {"word2idx": {...}, "tok_emb": ndarray, "dim": int}        # to_dict 직렬화형
      3) {"word2idx": {...}, "embeddings"/"emb"/"weight": ndarray}  # 키 이름 변형
      4) {"<단어>": ndarray, ...}                                    # 단어->벡터 딕셔너리
      5) gascore_engine 통째 dict 안에 위 키들이 중첩된 경우도 탐색
    self.load_error 에 실패 사유를 남겨 UI에서 표시할 수 있다.
    """
    EMB_KEYS = ("tok_emb", "embeddings", "embedding", "emb", "weight", "weights", "vectors")

    def __init__(self, dim=64, tok_emb_path: str | None = None):
        self.dim = dim
        self.mode = "hash"
        self.word2idx = None
        self.tok_emb = None
        self.load_error: str | None = None
        self.vocab_size = 0
        self._cache: dict[str, np.ndarray] = {}
        if tok_emb_path and os.path.exists(tok_emb_path):
            try:
                self._load_tok_emb(tok_emb_path)
            except Exception as e:
                self.load_error = f"{type(e).__name__}: {e}"

    @staticmethod
    def _find_emb_and_vocab(d: dict):
        """dict(중첩 포함)에서 임베딩 행렬과 word2idx를 찾아낸다."""
        # 1) 직접 키
        emb = None
        for k in EmbeddingProvider.EMB_KEYS:
            if k in d and d[k] is not None:
                emb = np.asarray(d[k])
                if emb.ndim == 2:
                    break
                emb = None
        w2i = d.get("word2idx") or d.get("vocab") or d.get("stoi")
        if emb is not None and w2i is not None:
            return emb, w2i
        # 2) idx2word만 있으면 뒤집어서 word2idx 생성
        if emb is not None and "idx2word" in d:
            i2w = d["idx2word"]
            if isinstance(i2w, dict):
                w2i = {v: int(k) for k, v in i2w.items()}
            else:  # list
                w2i = {w: i for i, w in enumerate(i2w)}
            return emb, w2i
        # 3) 한 단계 중첩 탐색 (예: {"engine": {...}} / {"som": ..., "tok": {...}})
        for v in d.values():
            if isinstance(v, dict):
                e, w = EmbeddingProvider._find_emb_and_vocab(v)
                if e is not None and w is not None:
                    return e, w
        return None, None

    def _load_tok_emb(self, path: str):
        with open(path, "rb") as f:
            d = pickle.load(f)

        # 형식 4: 단어->벡터 딕셔너리
        if isinstance(d, dict) and d and all(
            isinstance(k, str) for k in list(d.keys())[:20]
        ) and all(
            isinstance(v, (list, np.ndarray)) for v in list(d.values())[:5]
        ) and not any(k in d for k in ("word2idx", "tok_emb", "idx2word", "vocab")):
            words = list(d.keys())
            mat = np.asarray([np.asarray(d[w], dtype=np.float64) for w in words])
            self.word2idx = {w: i for i, w in enumerate(words)}
            self.tok_emb = mat
            self.dim = mat.shape[1]
            self.vocab_size = len(words)
            self.mode = "tok_emb"
            return

        if not isinstance(d, dict):
            self.load_error = "pkl 최상위가 dict가 아님 — 인식 불가"
            return

        emb, w2i = self._find_emb_and_vocab(d)
        if emb is None or w2i is None:
            self.load_error = ("tok_emb/word2idx를 찾지 못함. "
                               f"최상위 키: {list(d.keys())[:8]}")
            return

        self.tok_emb = np.asarray(emb, dtype=np.float64)
        # 정수 인덱스 보장
        self.word2idx = {str(k): int(v) for k, v in w2i.items()}
        self.dim = self.tok_emb.shape[1]
        self.vocab_size = self.tok_emb.shape[0]
        self.mode = "tok_emb"

    def apply_matrix(self, emb_matrix, word2idx):
        """코퍼스에서 생성한 tok_emb 행렬을 직접 적용 (자동 생성용)."""
        self.tok_emb = np.asarray(emb_matrix, dtype=np.float64)
        self.word2idx = {str(k): int(v) for k, v in word2idx.items()}
        self.dim = self.tok_emb.shape[1]
        self.vocab_size = self.tok_emb.shape[0]
        self.mode = "tok_emb"
        return self

    def _word_vec(self, w: str) -> np.ndarray:
        if self.mode == "tok_emb":
            idx = self.word2idx.get(w)
            if idx is not None and 0 <= idx < self.tok_emb.shape[0]:
                return self.tok_emb[idx]
            return np.zeros(self.dim)
        # 해시 폴백
        if w not in self._cache:
            h = int(hashlib.md5(w.encode("utf-8")).hexdigest(), 16)
            r = np.random.default_rng(h % (2**32))
            self._cache[w] = r.normal(size=self.dim)
        return self._cache[w]

    def encode(self, sentence: str) -> np.ndarray:
        words = sentence.split()
        if not words:
            return np.zeros(self.dim)
        vs = [self._word_vec(w) for w in words]
        return np.mean(vs, axis=0)

    def encode_many(self, sentences) -> np.ndarray:
        X = np.stack([self.encode(s) for s in sentences])
        X = (X - X.mean(0)) / (X.std(0) + 1e-8)
        return X


# ======================================================================
# 점진적으로 자라는 임베딩 (씨앗 + 자율 성장)
#   - 한국인이 모국어(씨앗)를 알고, 새 분야 전문어를 배우듯 동작
#   - 새 단어: vocab에 추가, 주변 단어 평균으로 초기화(콜드스타트 완화)
#   - 기존 단어: 새 문맥으로 skip-gram 미세조정
#   - 씨앗 단어는 천천히(보호), 새 단어는 빠르게 학습 → 망각 방지
# ======================================================================
class GrowingEmbedding:
    def __init__(self, dim=64, seed_lr=0.01, new_lr=0.05, window=3, neg=5, seed=0):
        self.dim = dim
        self.word2idx = {}
        self.idx2word = {}
        self.W = np.zeros((0, dim), dtype=np.float64)   # 중심 임베딩
        self.C = np.zeros((0, dim), dtype=np.float64)   # 문맥 임베딩
        self.freq = np.zeros(0, dtype=np.float64)
        self.is_seed = np.zeros(0, dtype=bool)          # 씨앗 단어 보호 플래그
        self.seed_lr = seed_lr                          # 씨앗 단어 학습률(작게)
        self.new_lr = new_lr                            # 새 단어 학습률(크게)
        self.window = window
        self.neg = neg
        self.rng = np.random.default_rng(seed)
        self.total_updates = 0

    # ---- 씨앗 적재: tok_emb pkl 또는 EmbeddingProvider에서 ----
    def load_seed(self, tok_emb_path=None, word2idx=None, matrix=None):
        if tok_emb_path:
            prov = EmbeddingProvider(dim=self.dim, tok_emb_path=tok_emb_path)
            if prov.mode != "tok_emb":
                return False, (prov.load_error or "tok_emb 인식 실패")
            word2idx, matrix = prov.word2idx, prov.tok_emb
        if word2idx is None or matrix is None:
            return False, "씨앗 데이터 없음"
        matrix = np.asarray(matrix, dtype=np.float64)
        self.dim = matrix.shape[1]
        V = matrix.shape[0]
        self.word2idx = dict(word2idx)
        self.idx2word = {i: w for w, i in self.word2idx.items()}
        self.W = matrix.copy()
        self.C = matrix.copy() * 0.1
        self.freq = np.ones(V)
        self.is_seed = np.ones(V, dtype=bool)
        return True, f"씨앗 {V}단어 적재(dim={self.dim})"

    def _add_word(self, w, init_vec=None):
        idx = len(self.word2idx)
        self.word2idx[w] = idx
        self.idx2word[idx] = w
        if init_vec is None:
            init_vec = (self.rng.random(self.dim) - 0.5) / self.dim
        self.W = np.vstack([self.W, init_vec])
        self.C = np.vstack([self.C, np.zeros(self.dim)])
        self.freq = np.append(self.freq, 0.0)
        self.is_seed = np.append(self.is_seed, False)
        return idx

    # ---- 새 문장들로 점진 학습 (vocab 성장 + 미세조정) ----
    def grow(self, sentences, epochs=3):
        # 1) 토크나이즈
        seqs = [tokenize(s) for s in sentences]
        # 2) 새 단어 등록 (주변 기존단어 평균으로 초기화 → 콜드스타트 완화)
        new_words = 0
        for toks in seqs:
            for i, w in enumerate(toks):
                if w not in self.word2idx:
                    ctx = [self.W[self.word2idx[t]] for t in toks
                           if t in self.word2idx]
                    init = np.mean(ctx, axis=0) if ctx else None
                    if init is not None:
                        init = init + self.rng.normal(scale=0.01, size=self.dim)
                    self._add_word(w, init)
                    new_words += 1
        # 3) 인덱스 시퀀스
        idx_seqs = [[self.word2idx[w] for w in toks if w in self.word2idx]
                    for toks in seqs]
        idx_seqs = [s for s in idx_seqs if len(s) >= 2]
        for s in idx_seqs:
            for i in s:
                self.freq[i] += 1
        # 4) negative sampling 분포
        p = self.freq ** 0.75
        p = p / p.sum() if p.sum() > 0 else None

        def sig(x):
            return 1.0 / (1.0 + np.exp(-np.clip(x, -30, 30)))

        for ep in range(epochs):
            self.rng.shuffle(idx_seqs)
            for s in idx_seqs:
                for i, center in enumerate(s):
                    win = self.rng.integers(1, self.window + 1)
                    lo, hi = max(0, i - win), min(len(s), i + win + 1)
                    for j in range(lo, hi):
                        if j == i:
                            continue
                        ctx = s[j]
                        negs = self.rng.choice(len(p), size=self.neg, p=p) if p is not None else []
                        targets = np.array([ctx] + list(negs))
                        labels = np.zeros(len(targets)); labels[0] = 1.0
                        v_in = self.W[center]
                        v_out = self.C[targets]
                        g = sig(v_out @ v_in) - labels
                        # 학습률: 중심단어가 씨앗이면 작게, 아니면 크게
                        lr_c = self.seed_lr if self.is_seed[center] else self.new_lr
                        grad_in = g @ v_out
                        # 문맥 업데이트(타깃별 씨앗 여부 반영)
                        lr_t = np.where(self.is_seed[targets], self.seed_lr, self.new_lr)
                        self.C[targets] -= (lr_t[:, None]) * np.outer(g, v_in)
                        self.W[center] -= lr_c * grad_in
                        self.total_updates += 1
        # 정규화는 encode 단계에서
        return {"new_words": new_words, "vocab": len(self.word2idx)}

    def _word_vec(self, w):
        idx = self.word2idx.get(w)
        if idx is not None:
            return self.W[idx]
        return np.zeros(self.dim)

    def encode(self, sentence):
        toks = tokenize(sentence)
        if not toks:
            return np.zeros(self.dim)
        vs = [self._word_vec(w) for w in toks]
        v = np.mean(vs, axis=0)
        return v

    def encode_many(self, sentences):
        X = np.stack([self.encode(s) for s in sentences])
        # 단위벡터 정규화 (SOM 붕괴 방지)
        X = X / (np.linalg.norm(X, axis=1, keepdims=True) + 1e-8)
        return X

    def save(self, path):
        import pickle as _pk
        _pk.dump({"word2idx": self.word2idx, "idx2word": self.idx2word,
                  "tok_emb": self.W.astype(np.float32), "dim": self.dim,
                  "is_seed": self.is_seed, "freq": self.freq}, open(path, "wb"))

    @property
    def vocab_size(self):
        return len(self.word2idx)


# ======================================================================
# 마르코프 도메인 가드레일
#   - 학습된 코퍼스의 토큰 바이그램 전이확률(JM smoothing)을 만든다
#   - 새 문장의 avg_logP를 채점 → 임계값 밖이면 '도메인 밖'으로 거부
#   - 자율 수집 시 "마르코프 거부권": 코퍼스에 넣기 전 문지기 역할
# ======================================================================
class MarkovGuardrail:
    def __init__(self, jm_lambda=0.7):
        self.uni = {}          # 유니그램 카운트
        self.bi = {}           # 바이그램 카운트 {(w1,w2): n}
        self.uni_total = 0
        self.vocab = set()
        self.jm = jm_lambda    # JM 보간 계수(바이그램 비중)
        self.trained = False

    def fit(self, sentences):
        self.__init__(self.jm)
        for s in sentences:
            toks = ["<s>"] + tokenize(s) + ["</s>"]
            for i, w in enumerate(toks):
                self.uni[w] = self.uni.get(w, 0) + 1
                self.uni_total += 1
                self.vocab.add(w)
                if i > 0:
                    key = (toks[i-1], w)
                    self.bi[key] = self.bi.get(key, 0) + 1
        self.trained = len(self.vocab) > 1
        return self

    def update(self, sentences):
        """기존 모델에 새 문장 누적(자율 학습 중 점진 갱신)."""
        for s in sentences:
            toks = ["<s>"] + tokenize(s) + ["</s>"]
            for i, w in enumerate(toks):
                self.uni[w] = self.uni.get(w, 0) + 1
                self.uni_total += 1
                self.vocab.add(w)
                if i > 0:
                    key = (toks[i-1], w)
                    self.bi[key] = self.bi.get(key, 0) + 1
        self.trained = len(self.vocab) > 1
        return self

    def _logP(self, w1, w2):
        V = max(1, len(self.vocab))
        # 유니그램 확률(add-1)
        p_uni = (self.uni.get(w2, 0) + 1) / (self.uni_total + V)
        # 바이그램 확률
        c_w1 = self.uni.get(w1, 0)
        p_bi = (self.bi.get((w1, w2), 0) / c_w1) if c_w1 > 0 else 0.0
        # Jelinek-Mercer 보간
        p = self.jm * p_bi + (1 - self.jm) * p_uni
        if p <= 0:
            p = 1e-10
        return math.log(p)

    def score(self, sentence):
        """문장의 평균 logP. 높을수록 도메인에 맞음(0에 가까움)."""
        if not self.trained:
            return 0.0
        toks = ["<s>"] + tokenize(sentence) + ["</s>"]
        if len(toks) < 2:
            return -99.0
        lp = sum(self._logP(toks[i-1], toks[i]) for i in range(1, len(toks)))
        return lp / (len(toks) - 1)

    def judge(self, sentence, threshold=-9.0):
        """threshold보다 낮으면 거부(도메인 밖). 반환: (통과여부, 점수)."""
        sc = self.score(sentence)
        return (sc >= threshold), sc

    def filter(self, sentences, threshold=-9.0):
        """문장 리스트를 채점해 통과/거부로 분리."""
        passed, rejected = [], []
        for s in sentences:
            ok, sc = self.judge(s, threshold)
            (passed if ok else rejected).append((s, sc))
        return passed, rejected

    def suggest_threshold(self, sentences, percentile=10, holdout=0.2):
        """
        임계값 추천. holdout(안 외운 도메인 문장)의 점수 분포에서 경계를 잡되,
        큰/단일도메인 코퍼스에서 과도하게 빡빡해지는 것을 막기 위해
        '중앙값 - 넉넉한 마진' 방식 + 절대 상/하한 클램프를 쓴다.
        """
        import random as _rnd
        if len(sentences) < 10:
            scores = sorted(self.score(s) for s in sentences)
            k = max(0, int(len(scores) * percentile / 100) - 1)
            return min(scores[k] if scores else -9.0, -6.0)
        sents = list(sentences)
        _rnd.Random(0).shuffle(sents)
        n_hold = max(10, int(len(sents) * holdout))
        hold = sents[:n_hold]
        train = sents[n_hold:]
        probe = MarkovGuardrail(self.jm).fit(train)
        scores = sorted(probe.score(s) for s in hold)
        if not scores:
            return -9.0
        median = scores[len(scores) // 2]
        low = scores[max(0, int(len(scores) * 0.1) - 1)]  # 하위 10%
        # 도메인 안 문장 대부분(중앙값 부근)을 통과시키도록,
        # 중앙값에서 (중앙값-하위10%) 만큼 더 내려 넉넉히 잡는다.
        spread = max(0.5, median - low)
        thr = low - spread * 1.0
        # 절대 클램프: 너무 빡빡(-1~-5)하거나 너무 느슨(-13)하지 않게
        thr = max(-12.0, min(thr, -6.0))
        return thr


# ======================================================================
# INT8 양자화
# ======================================================================
class Quantizer:
    def __init__(self, W):
        self.absmax = np.abs(W).max(axis=0) + 1e-8
        self.scale = self.absmax / 127.0

    def q(self, W):
        return np.clip(np.round(W / self.scale), -127, 127).astype(np.int8)

    def dq(self, Wq):
        return Wq.astype(np.float64) * self.scale


# ======================================================================
# Growing SOM
# ======================================================================
class GrowingSOM:
    def __init__(self, dim, init_nodes=16, grow_threshold=2.0, seed=0):
        rng = np.random.default_rng(seed)
        self.dim = dim
        self.W = rng.normal(scale=0.3, size=(init_nodes, dim))
        self.coords = rng.normal(scale=1.0, size=(init_nodes, 2))
        self.err = np.zeros(init_nodes)
        self.grow_threshold = grow_threshold
        self.round = 0

    @property
    def n(self):
        return self.W.shape[0]

    def bmu(self, x):
        d2 = np.einsum("nd,nd->n", self.W - x, self.W - x)
        i = int(np.argmin(d2))
        return i, float(np.sqrt(d2[i]))

    def train_step(self, X, lr, radius):
        for x in X:
            i, dist = self.bmu(x)
            self.err[i] += dist
            cd2 = np.einsum("nd,nd->n", self.coords - self.coords[i],
                            self.coords - self.coords[i])
            h = np.exp(-cd2 / (2 * radius ** 2))
            self.W += (lr * h)[:, None] * (x - self.W)

    def grow(self, max_add=4):
        added = 0
        order = np.argsort(self.err)[::-1]
        for i in order[:max_add]:
            if self.err[i] < self.grow_threshold:
                break
            new_w = self.W[i] + np.random.default_rng(self.n).normal(scale=0.1, size=self.dim)
            new_c = self.coords[i] + np.random.default_rng(self.n + 1).normal(scale=0.3, size=2)
            self.W = np.vstack([self.W, new_w])
            self.coords = np.vstack([self.coords, new_c])
            self.err = np.append(self.err, 0.0)
            self.err[i] *= 0.5
            added += 1
        return added

    # ---- 추론경로 (의미 가중 최단경로) ----
    def _knn_graph(self, k=6):
        D = np.linalg.norm(self.coords[:, None] - self.coords[None, :], axis=2)
        nbr = np.argsort(D, axis=1)[:, 1:k+1]
        return nbr

    def semantic_path(self, src, dst, k=6):
        import heapq
        if src == dst:
            return [src], 0.0, 0
        nbr = self._knn_graph(k)
        N = self.n
        dist = [float("inf")] * N
        prev = [-1] * N
        dist[src] = 0.0
        pq = [(0.0, src)]
        while pq:
            d, u = heapq.heappop(pq)
            if d > dist[u]:
                continue
            if u == dst:
                break
            for v in nbr[u]:
                w = np.linalg.norm(self.W[u] - self.W[v])
                nd = d + w
                if nd < dist[int(v)]:
                    dist[int(v)] = nd
                    prev[int(v)] = u
                    heapq.heappush(pq, (nd, int(v)))
        if not np.isfinite(dist[dst]):
            return [src, dst], float("inf"), -1
        path, cur = [], dst
        while cur != -1:
            path.append(cur); cur = prev[cur]
        path.reverse()
        return path, dist[dst], len(path) - 1


# ======================================================================
# 자율 검색어 생성 (GSOM 빈 격자 = 호기심)
#   - 입력을 BMU로 매핑해 노드별 밀도를 구한다
#   - 저밀도(덜 탐색된) 노드 근처의 대표 단어를 뽑아 검색어로 만든다
#   - 이렇게 만든 검색어로 크롤링하면 "안 배운 영역"을 채우게 된다
# ======================================================================
def learning_reward(qe_before, qe_after, vocab_before, vocab_after,
                    zpd_low=0.15, zpd_high=0.5,
                    w_assimilate=1.0, w_difficulty=0.6, w_diversity=0.4):
    """
    한 번의 학습(자료 투입+학습)이 '좋은 학습'이었는지 보상으로 평가.

    교육학적 직관:
      - assimilate: 결국 지도에 녹아들었는가 (QE가 내려갔는가) → 학습 성공
      - difficulty: 너무 쉽지도(QE 변화 거의 0) 너무 어렵지도(QE 폭증) 않은
                    '적당한 난이도'(ZPD, 근접발달영역)에 가산점
      - diversity : 어휘 다양성이 유지/증가하면 +, 단조 폭락이면 - (단, 모험 허용)

    반환: (총보상, 세부dict)
    """
    # 1) 동화(assimilation): QE가 내려갈수록 +. 상대 변화율 사용
    rel_qe = (qe_before - qe_after) / (abs(qe_before) + 1e-8)
    r_assim = w_assimilate * rel_qe          # QE 내려가면 +, 오르면 -

    # 2) 난이도(ZPD): '도전했고 결국 녹여낸' 경우만 보너스.
    #    QE가 올라가버린(동화 실패) 경우엔 난이도 보너스를 주지 않는다.
    challenge = abs(qe_after - qe_before) / (abs(qe_before) + 1e-8)
    assimilated = qe_after <= qe_before          # 결국 녹아들었는가
    if not assimilated:
        # 동화 실패(QE 상승) = 너무 어렵거나 잡음 → 난이도 보너스 없음, 감점
        over = min(2.0, challenge / zpd_high)
        r_diff = -w_difficulty * over * 0.5
    elif zpd_low <= challenge <= zpd_high:
        r_diff = w_difficulty * 1.0              # 적당히 도전 + 녹여냄 = 최고
    elif challenge < zpd_low:
        r_diff = w_difficulty * (challenge / zpd_low) * 0.5   # 너무 쉬움
    else:
        over = min(2.0, (challenge - zpd_high) / zpd_high)
        r_diff = w_difficulty * (1.0 - over)     # 녹였지만 좀 과함

    # 3) 다양성: 떨어져도 '모험'으로 어느 정도 허용 → 하락에 둔감, 상승에 민감
    d_vocab = vocab_after - vocab_before
    if d_vocab >= 0:
        r_div = w_diversity * d_vocab * 2.0           # 상승은 두 배로 보상
    else:
        # 하락은 모험으로 절반만 벌점 + 작은 하락은 무시(탐색 허용)
        tolerated = max(0.0, -d_vocab - 0.05)         # 0.05까지는 봐줌
        r_div = -w_diversity * tolerated * 0.5

    total = r_assim + r_diff + r_div
    return total, {
        "assimilate": round(r_assim, 3),
        "difficulty": round(r_diff, 3),
        "diversity": round(r_div, 3),
        "challenge": round(challenge, 3),
        "total": round(total, 3),
    }


class QueryPolicy:
    """
    보상 기반 자율 검색어 정책.
    - 각 검색어(또는 그 키워드)의 누적 보상을 기억
    - 다음 자율 탐색에서 보상 높았던 방향의 키워드에 가중치 부여
    - epsilon-greedy: 가끔은 새 방향도 탐색(exploration)
    교육학: '잘 배워지는 주제는 더 깊이, 가끔은 새 주제도 모험'
    """
    def __init__(self, epsilon=0.3):
        self.kw_reward = {}     # 키워드 -> 누적보상
        self.kw_count = {}      # 키워드 -> 시도횟수
        self.epsilon = epsilon
        self.log = []           # (검색어, 보상) 이력

    def record(self, query, reward):
        self.log.append((query, round(reward, 3)))
        for w in query.split():
            self.kw_reward[w] = self.kw_reward.get(w, 0.0) + reward
            self.kw_count[w] = self.kw_count.get(w, 0) + 1

    def score_keyword(self, w):
        """키워드의 평균 보상(미시도는 0=중립)."""
        if w not in self.kw_count:
            return 0.0
        return self.kw_reward[w] / self.kw_count[w]

    def rank_queries(self, candidate_queries, rng=None):
        """
        후보 검색어들을 보상 기대치로 정렬.
        epsilon 확률로는 무작위(탐색), 아니면 보상순(활용).
        """
        import random as _r
        rng = rng or _r.Random()
        if rng.random() < self.epsilon:
            shuffled = list(candidate_queries)
            rng.shuffle(shuffled)
            return shuffled, "explore"
        scored = sorted(candidate_queries,
                        key=lambda q: -np.mean([self.score_keyword(w)
                                                for w in q.split()] or [0]))
        return scored, "exploit"


def autonomous_queries(gsom, X, corpus, emb, n_queries=3, words_per_query=2):
    """
    반환: [검색어, ...]
    gsom  : GrowingSOM
    X     : 현재 코퍼스 임베딩 (encode_many 결과)
    corpus: 문장 리스트 (X와 같은 순서)
    emb   : EmbeddingProvider/GrowingEmbedding (단어 추출용)
    """
    if gsom.n == 0 or len(X) == 0:
        return []
    # 1) 각 입력의 BMU → 노드별 밀도
    bmus = []
    for x in X:
        d2 = np.einsum("nd,nd->n", gsom.W - x, gsom.W - x)
        bmus.append(int(np.argmin(d2)))
    density = np.zeros(gsom.n)
    for b in bmus:
        density[b] += 1

    # 2) 저밀도지만 '완전히 빈 건 아닌' 경계 노드 우선
    #    (완전 빈 노드는 의미 단서가 없으니, 1~2개 들어온 변두리를 탐색)
    cand = [i for i in range(gsom.n) if 0 < density[i] <= 2]
    if not cand:
        cand = list(np.argsort(density)[:n_queries * 3])

    # 3) 각 후보 노드에 매핑된 문장에서 핵심 단어 추출
    node_sents = {}
    for s, b in zip(corpus, bmus):
        node_sents.setdefault(b, []).append(s)

    # 단어 빈도(전역) — 흔한 단어 제외용
    from collections import Counter
    global_cnt = Counter()
    for s in corpus:
        global_cnt.update(tokenize(s))

    queries = []
    used = set()
    rng = np.random.default_rng(len(corpus))
    rng.shuffle(cand)
    # 검색어로 부적합한 흔한 조사/접미 패턴
    stop = {"그리고", "하지만", "그래서", "그러나", "또한", "위한", "통해", "대한",
            "이런", "저런", "그런", "있는", "없는", "하는", "되는", "같은"}
    def good_word(w):
        if len(w) < 2:
            return False
        if not re.search(r"[가-힣]", w):   # 영문 파편 제외(한글 우선)
            return len(w) >= 4 and w.isalpha()
        if w in stop:
            return False
        # 조사로 끝나는 어절 정리: 끝 1글자가 조사면 떼기
        return True
    def clean_word(w):
        # 흔한 조사 어미 제거(검색 친화적으로)
        for josa in ("으로", "에서", "에게", "과", "와", "을", "를", "이", "가",
                     "은", "는", "의", "도", "만", "과의", "와의", "들", "엔"):
            if w.endswith(josa) and len(w) - len(josa) >= 2:
                return w[:-len(josa)]
        return w
    for node in cand:
        sents = node_sents.get(node, [])
        if not sents:
            continue
        wc = Counter()
        for s in sents:
            for w in tokenize(s):
                if good_word(w) and global_cnt[w] <= len(corpus) * 0.3:
                    wc[clean_word(w)] += 1
        top = [w for w, _ in wc.most_common(words_per_query) if len(w) >= 2]
        if len(top) >= 1:
            q = " ".join(top)
            if q not in used and len(q) >= 2:
                used.add(q)
                queries.append(q)
        if len(queries) >= n_queries:
            break
    return queries


# ======================================================================
# 계측
# ======================================================================
def measure(gsom: GrowingSOM, X, tokens):
    bmus, qes, topo_fail = [], [], 0
    for x in X:
        d2 = np.einsum("nd,nd->n", gsom.W - x, gsom.W - x)
        order = np.argsort(d2)[:2]
        bmus.append(int(order[0]))
        qes.append(float(np.sqrt(d2[order[0]])))
        # 좌표상 1·2등 인접 여부
        c1, c2 = gsom.coords[order[0]], gsom.coords[order[1]]
        if np.linalg.norm(c1 - c2) > 1.5:
            topo_fail += 1
    occ = len(set(bmus))
    all_tok = [t for seq in tokens for t in seq]
    vocab = len(set(all_tok)) / len(all_tok) if all_tok else 0.0
    return {
        "nodes": gsom.n,
        "occupancy": occ,
        "occ_ratio": occ / gsom.n,
        "mean_qe": float(np.mean(qes)) if qes else 0.0,
        "topo_error": topo_fail / len(X) if len(X) else 0.0,
        "vocab_div": vocab,
    }


def collapse_warning(history, window=3):
    if len(history) < window + 1:
        return None
    rec = history[-(window+1):]
    def down(key):
        s = [h[key] for h in rec]
        return all(b <= a for a, b in zip(s[:-1], s[1:])) and s[0] > s[-1]
    if down("occupancy") and down("vocab_div") and down("mean_qe"):
        return "붕괴 경보: 점유율·어휘·QE 동반 감소 (자기출력 메아리방)"
    if down("occupancy") and rec[-1]["occ_ratio"] < 0.1:
        return "붕괴 경보: 점유율 10%↓ 한 점 수렴"
    # 어휘 다양성 급락(메아리/중복 폭증) — 더 넓은 창으로도 감지
    if len(history) >= 6:
        vd = [h["vocab_div"] for h in history]
        recent_max = max(vd[-6:])
        if vd[-1] < 0.12 and recent_max > 0.25:
            return ("붕괴 경보: 어휘 다양성 급락(중복·메아리 의심) — "
                    "수집 중복 제거를 확인하세요")
    # 점유율 만성 저하 + 노드 과성장
    if rec[-1]["occ_ratio"] < 0.15 and rec[-1]["nodes"] > 200:
        return "붕괴 경보: 점유율 15%↓ + 노드 과성장 (빈 격자 누적)"
    return None


# ======================================================================
# 검색·크롤러 (네트워크 필요 — 로컬 실행 시 작동)
# ======================================================================
def _clean_text(text: str) -> str:
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _is_garbage_sentence(s: str) -> bool:
    """PDF 바이너리/깨진 인코딩 문장 판별."""
    if not s:
        return True
    # PDF 내부 구조 흔적
    if any(k in s for k in ("endobj", "endstream", "ProcSet", "/PDF",
                            " obj ", "stream", "xref", "/Font", "/Type")):
        return True
    # 제어문자(널·비인쇄) 비율
    ctrl = sum(1 for ch in s if ord(ch) < 32 and ch not in "\t\n\r")
    if ctrl > 0:
        return True
    # '정상 문자' = 한글/기본영문/숫자/일반 문장부호/공백
    normal = len(re.findall(r"[가-힣A-Za-z0-9 .,!?\"'()%~\-:;]", s))
    if normal / len(s) < 0.75:
        return True
    # 한글 또는 (의미있는) 영문 단어가 최소한 있어야 함
    has_ko = bool(re.search(r"[가-힣]{2,}", s))
    has_en = bool(re.search(r"[A-Za-z]{3,}", s))
    if not (has_ko or has_en):
        return True
    # 의미 단어(2글자+ 한글어절, 3글자+ 영단어) 개수가 너무 적으면 파편
    ko_words = re.findall(r"[가-힣]{2,}", s)
    en_words = re.findall(r"[A-Za-z]{3,}", s)
    if len(ko_words) + len(en_words) < 2:
        return True
    # 짧은데 한글이 거의 없고 기호·대문자 뒤섞이면 PDF 파편
    if len(s) < 40 and not has_ko:
        upper = len(re.findall(r"[A-Z]", s))
        if upper / max(1, len(s)) > 0.2:
            return True
    return False


def _split_sentences(text: str, min_len=12, max_len=220):
    """한국어/영어가 섞인 웹문서를 학습용 문장으로 쪼갠다."""
    text = _clean_text(text)
    raw = re.split(r"(?<=[.!?。！？다요죠함음임됨됨니다])\s+|[\n\r]+|[•·]", text)
    out = []
    junk = ("cookie", "javascript", "copyright", "로그인", "회원가입", "개인정보", "구독", "광고")
    for s in raw:
        s = _clean_text(s)
        if not (min_len <= len(s) <= max_len):
            continue
        if any(j.lower() in s.lower() for j in junk):
            continue
        if _is_garbage_sentence(s):      # 깨진/PDF바이너리 문장 제거
            continue
        out.append(s)
    return out


# ======================================================================
# 로컬 파일(docx / pdf / txt)에서 코퍼스 추출
# ======================================================================
def extract_text_from_docx(path: str) -> str:
    """Word(.docx) 본문 + 표 셀 텍스트를 추출."""
    try:
        import docx
    except Exception:
        raise RuntimeError("python-docx 미설치: pip install python-docx")
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    # 표 안의 텍스트도 수집
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)


def extract_text_from_pdf(path: str) -> str:
    """PDF 본문 추출. pdfplumber 우선(레이아웃 양호), 실패 시 pypdf."""
    text = ""
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(path) as pdf:
            for pg in pdf.pages:
                t = pg.extract_text() or ""
                if t.strip():
                    pages.append(t)
        text = "\n".join(pages)
    except Exception:
        text = ""
    if not text.strip():
        try:
            from pypdf import PdfReader
            reader = PdfReader(path)
            text = "\n".join((pg.extract_text() or "") for pg in reader.pages)
        except Exception as e:
            raise RuntimeError(f"PDF 추출 실패: {e}")
    return text


def extract_corpus_from_file(path: str, filename: str | None = None,
                             min_len=12, max_len=220):
    """
    업로드 파일 경로에서 학습용 문장 리스트를 추출한다.
    지원: .docx .pdf .txt .md
    반환: (sentences, info) — info는 진단 문자열
    """
    name = (filename or path).lower()
    try:
        if name.endswith(".docx"):
            raw = extract_text_from_docx(path)
            kind = "docx"
        elif name.endswith(".pdf"):
            raw = extract_text_from_pdf(path)
            kind = "pdf"
        elif name.endswith((".txt", ".md")):
            with open(path, "rb") as f:
                data = f.read()
            raw = None
            for enc in ("utf-8", "cp949", "euc-kr", "latin-1"):
                try:
                    raw = data.decode(enc); break
                except Exception:
                    continue
            raw = raw or data.decode("utf-8", "ignore")
            kind = "txt"
        elif name.endswith(".doc"):
            return [], "구버전 .doc는 미지원입니다. .docx로 저장 후 올려주세요."
        else:
            return [], f"지원하지 않는 형식: {name.split('.')[-1]}"
    except Exception as e:
        return [], f"{kind if 'kind' in dir() else '파일'} 추출 오류: {e}"

    sents = _split_sentences(raw, min_len=min_len, max_len=max_len)
    # 중복 제거(순서 보존)
    seen, uniq = set(), []
    for s in sents:
        if s not in seen:
            seen.add(s); uniq.append(s)
    info = f"{kind} · 원문 {len(raw)}자 → 문장 {len(uniq)}개"
    return uniq, info


_BROWSER_HEADERS = {
    "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                   "AppleWebKit/537.36 (KHTML, like Gecko) "
                   "Chrome/124.0.0.0 Safari/537.36"),
    "Accept": ("text/html,application/xhtml+xml,application/xml;q=0.9,"
               "image/avif,image/webp,*/*;q=0.8"),
    "Accept-Language": "ko-KR,ko;q=0.9,en-US;q=0.8,en;q=0.7",
    "Accept-Encoding": "identity",
    "Referer": "https://www.google.com/",
    "Connection": "keep-alive",
    "Upgrade-Insecure-Requests": "1",
}


def _fetch_url_bytes(url: str, timeout=12):
    """URL에서 raw bytes와 content-type을 가져온다."""
    import urllib.request
    req = urllib.request.Request(url, headers=_BROWSER_HEADERS)
    with urllib.request.urlopen(req, timeout=timeout) as r:
        raw = r.read()
        ctype = r.headers.get("Content-Type", "")
    return raw, ctype


def _pdf_bytes_to_text(raw: bytes) -> str:
    """PDF 바이트에서 텍스트 추출 (pdfplumber 우선, pypdf 폴백)."""
    import io as _io
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(_io.BytesIO(raw)) as pdf:
            parts = []
            for pg in pdf.pages:
                t = pg.extract_text() or ""
                if t.strip():
                    parts.append(t)
            text = "\n".join(parts)
    except Exception:
        text = ""
    if not text.strip():
        try:
            from pypdf import PdfReader
            reader = PdfReader(_io.BytesIO(raw))
            text = "\n".join((pg.extract_text() or "") for pg in reader.pages)
        except Exception:
            text = ""
    return text


def _fetch_url_text(url: str, timeout=12) -> str:
    raw, ctype = _fetch_url_bytes(url, timeout=timeout)
    # PDF 감지: content-type 또는 URL 확장자 또는 매직넘버(%PDF)
    is_pdf = ("application/pdf" in ctype.lower()
              or url.lower().split("?")[0].endswith(".pdf")
              or raw[:5] == b"%PDF-")
    if is_pdf:
        return _pdf_bytes_to_text(raw)
    # HTML/텍스트 디코딩
    for enc in ("utf-8", "cp949", "euc-kr", "latin-1"):
        try:
            return raw.decode(enc, "ignore")
        except Exception:
            pass
    return raw.decode("utf-8", "ignore")


class _TextExtractorHTML:
    """html.parser를 함수 내부 import 없이 쓰기 위한 가벼운 본문 추출기."""
    pass


def _html_to_text(html: str) -> str:
    from html.parser import HTMLParser

    class TextExtractor(HTMLParser):
        def __init__(self):
            super().__init__()
            self.skip_depth = 0
            self.parts = []
        def handle_starttag(self, tag, attrs):
            if tag.lower() in ("script", "style", "noscript", "svg", "canvas", "header", "footer", "nav"):
                self.skip_depth += 1
        def handle_endtag(self, tag):
            if tag.lower() in ("script", "style", "noscript", "svg", "canvas", "header", "footer", "nav") and self.skip_depth:
                self.skip_depth -= 1
        def handle_data(self, data):
            if self.skip_depth == 0:
                t = data.strip()
                if t:
                    self.parts.append(t)

    ex = TextExtractor()
    ex.feed(html)
    return " ".join(ex.parts)


def _extract_links_generic(html: str, max_results: int):
    """검색결과 HTML에서 외부 URL 추출 (DDG 리다이렉트 디코딩 포함)."""
    import urllib.parse
    links = []
    for m in re.findall(r'href=["\']([^"\']+)["\']', html):
        href = m.replace("&amp;", "&")
        real = None
        if "uddg=" in href:  # DuckDuckGo 리다이렉트
            qs = urllib.parse.parse_qs(urllib.parse.urlparse(href).query)
            if "uddg" in qs:
                real = urllib.parse.unquote(qs["uddg"][0])
        elif "/url?q=" in href:  # Google 리다이렉트
            qs = urllib.parse.parse_qs(urllib.parse.urlparse(href).query)
            if "q" in qs:
                real = qs["q"][0]
        elif href.startswith("http://") or href.startswith("https://"):
            real = href
        if not real:
            continue
        if any(b in real for b in ["duckduckgo.com", "bing.com", "google.com",
                                   "javascript:", "mailto:", "microsoft.com/",
                                   "go.microsoft.com", "w3.org"]):
            continue
        if real not in links:
            links.append(real)
        if len(links) >= max_results:
            break
    return links


def brave_search(query: str, api_key: str, count=10, country="kr",
                 search_lang="ko", freshness=None, timeout=15):
    """
    Brave Search API (web/search)로 검색.
    반환: (results, error) — results는 [{title,url,description,extra}], error는 실패시 str.
    """
    import json
    import urllib.request, urllib.parse
    if not api_key:
        return [], "Brave API Key 없음"
    params = {
        "q": query[:380],
        "count": min(int(count), 20),
        "country": country,
        "search_lang": search_lang,
        "extra_snippets": "true",
    }
    if freshness:
        params["freshness"] = freshness  # pd/pw/pm/py
    url = "https://api.search.brave.com/res/v1/web/search?" + urllib.parse.urlencode(params)
    req = urllib.request.Request(url, headers={
        "Accept": "application/json",
        "Accept-Encoding": "identity",
        "X-Subscription-Token": api_key,
    })
    try:
        with urllib.request.urlopen(req, timeout=timeout) as r:
            data = json.loads(r.read().decode("utf-8", "ignore"))
    except urllib.error.HTTPError as e:
        code = e.code
        hint = {401: "API Key가 틀렸거나 누락(401)",
                422: "쿼리 파라미터 오류(422)",
                429: "사용량/요율 초과(429) — 잠시 후 재시도"}.get(code, f"HTTP {code}")
        return [], f"Brave 검색 실패: {hint}"
    except Exception as e:
        return [], f"Brave 검색 실패: {type(e).__name__}: {e}"

    out = []
    for item in (data.get("web", {}).get("results") or []):
        out.append({
            "title": item.get("title", ""),
            "url": item.get("url", ""),
            "description": item.get("description", ""),
            "extra": item.get("extra_snippets") or [],
        })
    return out, None


def brave_collect(query: str, api_key: str, max_pages=8, fetch_bodies=True,
                  max_sentences=300, country="kr", search_lang="ko",
                  freshness=None, delay=0.3):
    """
    Brave로 검색 → URL 수집 → (옵션) 본문 크롤링까지 해서 학습 문장 추출.

    fetch_bodies=True : 검색된 URL을 직접 크롤링(가장 알찬 코퍼스)
    fetch_bodies=False: 검색 스니펫/extra_snippets만으로 코퍼스 구성(빠름, API만으로 완결)

    반환: (sentences, sources, links)
    """
    results, err = brave_search(query, api_key, count=max_pages,
                                country=country, search_lang=search_lang,
                                freshness=freshness)
    if err:
        return [], [err], []
    if not results:
        return [], ["Brave 검색 결과 0건"], []

    links = [r["url"] for r in results if r["url"]]

    if not fetch_bodies:
        # 스니펫만으로 코퍼스 구성
        sents = []
        for r in results:
            for chunk in [r["description"]] + list(r["extra"]):
                sents.extend(_split_sentences(chunk or ""))
        seen, uniq = set(), []
        for s in sents:
            if s not in seen:
                seen.add(s); uniq.append(s)
        src = [f"Brave 스니펫 {len(results)}건 → 문장 {len(uniq)}개"]
        return uniq[:max_sentences], src, links

    # 본문 크롤링 (검색은 Brave가, 본문은 기존 크롤러가)
    sentences, sources = crawl_urls(links[:max_pages],
                                    max_sentences=max_sentences, delay=delay)
    sources = [f"Brave 검색 {len(links)}개 URL"] + sources
    return sentences, sources, links[:max_pages]


def _search_links(query: str, max_results=8):
    """
    여러 검색 엔진을 순서대로 시도. 하나가 막히면 다음으로 폴백.
    반환: (links, log) — log는 각 엔진 시도 결과(디버그/화면표시용).
    """
    import urllib.parse
    q = urllib.parse.quote(query)
    engines = [
        ("DuckDuckGo Lite", f"https://lite.duckduckgo.com/lite/?q={q}"),
        ("DuckDuckGo HTML", f"https://html.duckduckgo.com/html/?q={q}"),
        ("Bing", f"https://www.bing.com/search?q={q}&setlang=ko"),
        ("Mojeek", f"https://www.mojeek.com/search?q={q}"),
    ]
    log = []
    for name, url in engines:
        try:
            html = _fetch_url_text(url)
            links = _extract_links_generic(html, max_results)
            log.append(f"{name}: {len(links)}개 링크")
            if links:
                return links, log
        except Exception as e:
            log.append(f"{name}: 실패({type(e).__name__} {e})")
            continue
    return [], log


def _duckduckgo_links(query: str, max_results=8):
    """하위호환 래퍼."""
    links, _ = _search_links(query, max_results)
    return links


def crawl_urls(urls, max_sentences=300, delay=0.3):
    """URL 목록을 직접 크롤링해서 문장 리스트와 소스 로그를 반환한다."""
    sentences, sources = [], []
    for url in urls:
        try:
            html = _fetch_url_text(url)
            text = _html_to_text(html)
            ss = _split_sentences(text)
            if ss:
                sources.append({"url": url, "sentences": len(ss)})
                sentences.extend(ss)
            if delay:
                time.sleep(delay)
        except Exception as e:
            sources.append({"url": url, "error": f"{type(e).__name__}: {e}"})
            continue
    # 중복 제거
    seen, out = set(), []
    for s in sentences:
        key = s.lower().strip()
        if key not in seen:
            seen.add(key); out.append(s)
        if len(out) >= max_sentences:
            break
    return out, sources


def crawl_topic(topic: str, max_pages=5, max_sentences=300, extra_urls=None,
                delay=0.3, return_sources=False,
                brave_api_key=None, brave_fetch_bodies=True,
                country="kr", search_lang="ko", freshness=None):
    """
    '이 분야 검색해서 학습' 명령 처리.

    검색 경로 우선순위:
      1) Brave API 키가 있으면 Brave Search 사용 (안정적, 권장)
      2) 키 없으면 무료 검색엔진 스크래핑 폴백 (DDG/Bing/Mojeek — 자주 차단됨)
    사용자가 직접 넣은 URL(extra_urls)은 항상 함께 크롤링.

    return_sources=True면 (sentences, sources, links)를 반환한다.
    """
    if not topic and not extra_urls:
        return ([], ["입력된 검색어/URL 없음"], []) if return_sources else []

    # ---- 경로 1: Brave API ----
    if topic and topic.strip() and brave_api_key:
        sents, sources, links = brave_collect(
            topic.strip(), brave_api_key, max_pages=max_pages,
            fetch_bodies=brave_fetch_bodies, max_sentences=max_sentences,
            country=country, search_lang=search_lang, freshness=freshness, delay=delay)
        # 사용자가 직접 URL도 넣었으면 추가 크롤링
        if extra_urls:
            extra = [u.strip() for u in extra_urls
                     if u.strip().startswith(("http://", "https://"))]
            if extra:
                es, esrc = crawl_urls(extra, max_sentences=max_sentences, delay=delay)
                sents = sents + es
                sources = sources + esrc
                links = links + extra
        if return_sources:
            return sents, sources, links
        return sents

    # ---- 경로 2: 스크래핑 폴백 ----
    urls = []
    search_log = []
    if topic and topic.strip():
        found, search_log = _search_links(topic.strip(), max_results=max_pages)
        urls.extend(found)
    if extra_urls:
        for u in extra_urls:
            u = u.strip()
            if u and (u.startswith("http://") or u.startswith("https://")) and u not in urls:
                urls.append(u)

    if not urls:
        diag = ["검색 결과 0건 — 엔진별 시도:"] + search_log
        diag.append("→ 무료 검색엔진이 모두 차단되었습니다. "
                    "Brave API 키를 넣거나 URL을 직접 입력하세요.")
        return ([], diag, []) if return_sources else []

    sentences, sources = crawl_urls(urls[:max_pages], max_sentences=max_sentences, delay=delay)
    if return_sources:
        return sentences, sources + search_log, urls[:max_pages]
    return sentences


# ======================================================================
# LLM 브리지 (답변 루프 — API 키 필요)
# ======================================================================
def llm_answer(
    question: str,
    context_sentences,
    model="gpt-4o-mini",
    temperature=0.3,
    max_tokens=1000,
    api_key: str | None = None,
    base_url: str | None = None,
    system_prompt: str | None = None,
):
    """
    학습된 SOM 코퍼스에서 뽑은 맥락으로 OpenAI-compatible LLM 답변 생성.

    - api_key가 전달되면 우선 사용하고, 없으면 OPENAI_API_KEY 환경변수를 사용합니다.
    - base_url을 전달하면 OpenAI 호환 서버(Ollama proxy, vLLM, LM Studio 등)에도 연결할 수 있습니다.
      예: https://api.openai.com/v1 또는 http://localhost:1234/v1
    """
    try:
        from openai import OpenAI
    except Exception:
        return "[openai 미설치] pip install openai 후 사용"

    key = api_key or os.environ.get("OPENAI_API_KEY")
    if not key:
        return "[API KEY 미설정] 사이드바에 API Key를 입력하거나 OPENAI_API_KEY 환경변수를 설정하세요."

    kwargs = {"api_key": key}
    if base_url:
        kwargs["base_url"] = base_url.rstrip("/")
    client = OpenAI(**kwargs)

    ctx = "\n".join(f"- {s}" for s in context_sentences[:15])
    user = f"[학습된 맥락]\n{ctx}\n\n[질문]\n{question}"

    default_system = (
        "당신은 주어진 문서 맥락에 근거해 답하는 전문 어시스턴트입니다.\n"
        "규칙:\n"
        "1. 아래 맥락을 최우선 근거로 사용하세요.\n"
        "2. 맥락에 있는 구체적 내용(수치, 용어, 절차)을 활용해 상세히 답하세요.\n"
        "3. 맥락에 없는 내용은 일반 지식으로 보완하되, 맥락과 모순되면 안 됩니다.\n"
        "4. 한국어로, 충분히 구체적이고 친절하게 답하세요."
    )
    messages = [{"role": "system", "content": system_prompt or default_system}]
    messages.append({"role": "user", "content": user})

    try:
        resp = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=temperature,
            max_tokens=max_tokens,
        )
        return resp.choices[0].message.content
    except Exception as e:
        return f"[LLM 호출 실패] {type(e).__name__}: {e}"


# ======================================================================
# 전체 상태 + pkl 누적 저장/불러오기
# ======================================================================
class SelfLoopState:
    def __init__(self, dim=64):
        self.dim = dim
        self.gsom = GrowingSOM(dim=dim)
        self.corpus: list[str] = []       # 학습된 모든 문장(코퍼스)
        self.history: list[dict] = []     # 라운드별 계측
        self.guardrail = MarkovGuardrail() # 도메인 마르코프 가드레일
        self.guard_threshold = -9.5        # 거부 임계값
        self.rejected_log: list = []       # 거부된 문장 기록(검토용)
        self.policy = QueryPolicy(epsilon=0.3)  # 보상 기반 검색어 정책
        self.created = time.time()

    def fit_guardrail(self, percentile=15):
        """현재 코퍼스로 가드레일 학습 + 임계값 자동 추천."""
        if len(self.corpus) >= 10:
            self.guardrail.fit(self.corpus)
            self.guard_threshold = self.guardrail.suggest_threshold(
                self.corpus, percentile=percentile)
        return self.guard_threshold

    def add_sentences(self, sentences, use_guardrail=True):
        """
        코퍼스에 문장 추가. use_guardrail=True면 가드레일 통과분만 저장.
        중복(이미 코퍼스에 있는 문장)은 자동 제거 → '학습 메아리 붕괴' 방지.
        반환: (추가된 수, 거부된 수)
        """
        # --- 중복 제거: 기존 코퍼스 + 이번 입력 내부 중복 모두 ---
        existing = set(self.corpus)
        deduped = []
        seen = set()
        for s in sentences:
            if s in existing or s in seen:
                continue
            seen.add(s)
            deduped.append(s)
        sentences = deduped
        if not sentences:
            return 0, 0

        if not use_guardrail or not self.guardrail.trained:
            self.corpus += sentences
            return len(sentences), 0
        passed, rejected = self.guardrail.filter(sentences, self.guard_threshold)
        kept = [s for s, _ in passed]
        self.corpus += kept
        self.rejected_log += [(s, round(sc, 2)) for s, sc in rejected][-200:]
        return len(kept), len(rejected)

    def save(self, path):
        # SOM 가중치는 INT8 양자화해서 저장 (용량 보완)
        quant = Quantizer(self.gsom.W)
        blob = {
            "dim": self.dim,
            "W_int8": quant.q(self.gsom.W),
            "W_scale": quant.scale,
            "coords": self.gsom.coords,
            "err": self.gsom.err,
            "round": self.gsom.round,
            "corpus": self.corpus,
            "history": self.history,
            "created": self.created,
            "policy_reward": self.policy.kw_reward,
            "policy_count": self.policy.kw_count,
            "policy_log": self.policy.log[-300:],
        }
        with open(path, "wb") as f:
            pickle.dump(blob, f)

    @classmethod
    def load(cls, path):
        with open(path, "rb") as f:
            b = pickle.load(f)
        st = cls(dim=b["dim"])
        # 역양자화로 SOM 복원
        st.gsom.W = b["W_int8"].astype(np.float64) * b["W_scale"]
        st.gsom.coords = b["coords"]
        st.gsom.err = b["err"]
        st.gsom.round = b.get("round", 0)
        st.corpus = b["corpus"]
        st.history = b["history"]
        st.created = b.get("created", time.time())
        # 정책 복원(옛 pkl엔 없을 수 있음)
        st.policy.kw_reward = b.get("policy_reward", {})
        st.policy.kw_count = b.get("policy_count", {})
        st.policy.log = b.get("policy_log", [])
        # 코퍼스로 가드레일 재구성 (옛 pkl도 호환)
        if len(st.corpus) >= 10:
            try:
                st.fit_guardrail()
            except Exception:
                pass
        return st


# ======================================================================
# 코퍼스에서 tok_emb 자동 생성 (제품용: 사용자가 tok_emb 신경 안 쓰게)
# ======================================================================
def build_tok_emb_from_corpus(corpus, dim=64, min_count=2, epochs=10,
                              window=4, neg=5, lr=0.05, seed=0,
                              max_sentences=4000):
    """
    코퍼스(문장 리스트)로 skip-gram tok_emb를 즉석 생성.
    반환: (tok_emb 행렬, word2idx) 또는 (None, None) if 자료 부족.
    자료가 적으면 품질이 낮으므로 호출측에서 분량을 확인할 것.
    """
    from collections import Counter
    sents = [s for s in corpus if s and not _is_garbage_sentence(s)]
    if len(sents) > max_sentences:
        # 너무 많으면 샘플링(속도)
        import random as _r
        _r.Random(seed).shuffle(sents)
        sents = sents[:max_sentences]
    cnt = Counter()
    for s in sents:
        cnt.update(tokenize(s))
    # 자료 적으면 min_count 완화
    mc = min_count if len(sents) >= 200 else 1
    vocab = [w for w, c in cnt.items() if c >= mc]
    if len(vocab) < 10:
        return None, None
    w2i = {w: i for i, w in enumerate(vocab)}
    V = len(w2i)
    rng = np.random.default_rng(seed)
    Win = (rng.random((V, dim)) - 0.5) / dim
    Wout = np.zeros((V, dim))
    freq = np.array([cnt[w] for w in vocab], float)
    pneg = freq ** 0.75
    pneg /= pneg.sum()
    seqs = [[w2i[w] for w in tokenize(s) if w in w2i] for s in sents]
    seqs = [s for s in seqs if len(s) >= 2]
    if not seqs:
        return None, None

    def sig(x):
        return 1.0 / (1.0 + np.exp(-np.clip(x, -30, 30)))

    for ep in range(epochs):
        rng.shuffle(seqs)
        for s in seqs:
            for i, ctr in enumerate(s):
                lo, hi = max(0, i - window), min(len(s), i + window + 1)
                for j in range(lo, hi):
                    if j == i:
                        continue
                    negs = rng.choice(V, neg, p=pneg)
                    tg = np.concatenate(([s[j]], negs))
                    lb = np.zeros(neg + 1)
                    lb[0] = 1.0
                    vin = Win[ctr]
                    vout = Wout[tg]
                    g = sig(vout @ vin) - lb
                    Wout[tg] -= lr * np.outer(g, vin)
                    Win[ctr] -= lr * (g @ vout)
    norms = np.linalg.norm(Win, axis=1, keepdims=True) + 1e-8
    return (Win / norms).astype(np.float32), w2i


# ======================================================================
# TF-IDF 검색 엔진 (옛 GasCore 에서 이식 — 문서 기반 정확한 맥락 추출)
# tok_emb 유사도가 부정확할 때도, 단어 겹침으로 관련 문장을 확실히 찾음.
# ======================================================================
# 검색 전용 조사/어미 분리 (핵심어 매칭률 향상: '청구항이'→'청구항')
_SEARCH_JOSA = ["으로부터", "에서부터", "에게서", "으로서", "이라고", "으로는",
                "에서는", "에게는", "에서", "에게", "한테", "으로", "이랑",
                "부터", "까지", "마다", "보다", "처럼", "만큼", "이나",
                "에도", "에만", "에는", "이라", "이고", "이가", "이를", "이는",
                "와", "과", "랑", "로", "을", "를", "은", "는", "이", "가",
                "의", "도", "만", "에"]

def _search_tokenize(sentence: str):
    """검색 전용: 기본 토큰화 후 조사를 한 번 더 떼어 핵심어 매칭률을 높인다."""
    base = tokenize(sentence)
    out = []
    for t in base:
        if re.fullmatch(r"[가-힣]+", t):
            for j in sorted(_SEARCH_JOSA, key=len, reverse=True):
                if t.endswith(j) and len(t) > len(j) + 1:
                    t = t[:-len(j)]
                    break
        out.append(t)
    return out


class CorpusSearcher:
    def __init__(self):
        self.sentences = []
        self.vocab = {}
        self.idf = {}
        self.vecs = None
        self._emb_cache = None
        self._emb_cache_n = 0

    def build(self, sentences):
        """문장 리스트로 TF-IDF 인덱스 구축."""
        import numpy as _np
        from collections import Counter as _C
        sents = list(dict.fromkeys(s.strip() for s in sentences
                                   if s and len(s.strip()) >= 6))
        self.sentences = sents
        if not sents:
            self.vecs = None
            return self
        all_toks = []
        for s in sents:
            all_toks += _search_tokenize(s)
        cnt = _C(all_toks)
        self.vocab = {w: i for i, w in enumerate(w for w, c in cnt.most_common() if c >= 1)}
        V = len(self.vocab)
        N = len(sents)
        df = _C()
        for s in sents:
            for t in set(_search_tokenize(s)):
                if t in self.vocab:
                    df[t] += 1
        self.idf = {w: _np.log((N + 1) / (df.get(w, 0) + 1)) + 1 for w in self.vocab}

        def tv(text):
            v = _np.zeros(V)
            toks = _search_tokenize(text)
            ct = _C(toks)
            for w, c in ct.items():
                if w in self.vocab:
                    v[self.vocab[w]] = (c / max(len(toks), 1)) * self.idf.get(w, 1.0)
            nrm = _np.linalg.norm(v)
            return v / nrm if nrm > 1e-12 else v

        self.vecs = _np.array([tv(s) for s in sents])
        return self

    def search(self, query, topk=5):
        """질문과 가장 유사한 문장 topk 반환: [(문장, 유사도), ...]"""
        import numpy as _np
        from collections import Counter as _C
        if not self.sentences or self.vecs is None:
            return []
        V = len(self.vocab)
        v = _np.zeros(V)
        toks = _search_tokenize(query)
        ct = _C(toks)
        for w, c in ct.items():
            if w in self.vocab:
                v[self.vocab[w]] = (c / max(len(toks), 1)) * self.idf.get(w, 1.0)
        nrm = _np.linalg.norm(v)
        if nrm < 1e-12:
            return []
        v /= nrm
        sims = self.vecs @ v
        idx = _np.argsort(-sims)[:topk]
        return [(self.sentences[i], float(sims[i])) for i in idx if sims[i] > 0.01]

    def search_hybrid(self, query, emb=None, topk=5, w_tfidf=0.5, w_emb=0.5):
        """
        하이브리드 검색: TF-IDF(단어 겹침) + tok_emb(의미) 결합.
        - TF-IDF는 드문 핵심어(특허/청구항)를 잡고
        - tok_emb는 의미적 유사성을 잡아
        흔한 표현('구성되어 있나')에 휘둘리는 문제를 완화한다.
        emb 가 hash 모드면 의미 점수가 약하므로 TF-IDF 비중을 자동으로 올린다.
        """
        import numpy as _np
        from collections import Counter as _C
        if not self.sentences or self.vecs is None:
            return []
        # --- TF-IDF 점수 ---
        V = len(self.vocab)
        v = _np.zeros(V)
        toks = _search_tokenize(query)
        ct = _C(toks)
        for w, c in ct.items():
            if w in self.vocab:
                v[self.vocab[w]] = (c / max(len(toks), 1)) * self.idf.get(w, 1.0)
        nrm = _np.linalg.norm(v)
        tfidf_sims = (self.vecs @ (v / nrm)) if nrm > 1e-12 else _np.zeros(len(self.sentences))

        # --- tok_emb 의미 점수 ---
        emb_sims = _np.zeros(len(self.sentences))
        use_emb = emb is not None and getattr(emb, "mode", "hash") == "tok_emb"
        if use_emb:
            qv = emb.encode(query)
            qn = qv / (_np.linalg.norm(qv) + 1e-12)
            if self._emb_cache is None or self._emb_cache_n != len(self.sentences):
                M = _np.array([emb.encode(s) for s in self.sentences])
                Mn = M / (_np.linalg.norm(M, axis=1, keepdims=True) + 1e-12)
                self._emb_cache = Mn
                self._emb_cache_n = len(self.sentences)
            emb_sims = self._emb_cache @ qn
        else:
            # 의미 점수 못 쓰면 TF-IDF에 전적으로 의존
            w_tfidf, w_emb = 1.0, 0.0

        # --- 정규화 후 결합 ---
        def _norm01(a):
            lo, hi = a.min(), a.max()
            return (a - lo) / (hi - lo) if hi - lo > 1e-12 else a * 0
        score = w_tfidf * _norm01(tfidf_sims) + w_emb * _norm01(emb_sims)
        idx = _np.argsort(-score)[:topk]
        return [(self.sentences[i], float(score[i])) for i in idx if score[i] > 0.01]
